from aiogram import Bot, Dispatcher, types, executor
from aiogram.types.message import ContentType
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from config import api_token, sponsor_channel_link, sponsor_channel_id
import shutil
import os
import logging
from PIL import Image
from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfMerger

dir_path = os.path.dirname(os.path.realpath(__file__))

log_file_path = dir_path + '/pdfbot.log'
logging.basicConfig(
    filename=log_file_path,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

bot = Bot(token=api_token)
storage = MemoryStorage()
dp = Dispatcher(bot, storage=storage)

user_modes = {}
photos_id = {}
docs_id = {}
file_names = {}

class MergeState(StatesGroup):
    waiting_for_filename = State()

class ConvertState(StatesGroup):
    waiting_for_filename = State()

def main_menu_keyboard():
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    keyboard.add('Изображения в PDF/DOCX')
    keyboard.add('Соединение PDF/DOCX в один')
    return keyboard

def back_keyboard():
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    keyboard.add('Назад')
    return keyboard

@dp.message_handler(commands='start')
async def start_handler(message: types.Message):
    user_id = str(message.chat.id)
    user_modes[user_id] = None
    photos_id.pop(user_id, None)
    docs_id.pop(user_id, None)
    file_names.pop(user_id, None)
    await message.answer('Выберите режим работы бота:', reply_markup=main_menu_keyboard())

@dp.message_handler(lambda message: message.text == 'Назад')
async def back_handler(message: types.Message, state: FSMContext):
    user_id = str(message.chat.id)
    user_modes[user_id] = None
    photos_id.pop(user_id, None)
    docs_id.pop(user_id, None)
    file_names.pop(user_id, None)
    await state.finish()
    await message.answer('Вы вернулись в главное меню. Выберите режим работы бота:', reply_markup=main_menu_keyboard())

@dp.message_handler(lambda message: message.text in ['Изображения в PDF/DOCX', 'Соединение PDF/DOCX в один'])
async def mode_select_handler(message: types.Message):
    user_id = str(message.chat.id)
    photos_id.pop(user_id, None)
    docs_id.pop(user_id, None)
    file_names.pop(user_id, None)

    if message.text == 'Изображения в PDF/DOCX':
        user_modes[user_id] = 'mode_1'
        await message.answer('Режим: Изображения в PDF/DOCX.\nОтправьте изображения для конвертации.', reply_markup=back_keyboard())
    else:
        user_modes[user_id] = 'mode_2'
        await message.answer('Режим: Соединение PDF/DOCX.\nОтправьте PDF или DOCX файлы для объединения.', reply_markup=back_keyboard())

@dp.message_handler(content_types=[ContentType.PHOTO, ContentType.DOCUMENT])
async def handle_files(message: types.Message, state: FSMContext):
    user_id = str(message.chat.id)
    mode = user_modes.get(user_id)

    if mode == 'mode_1':
        if message.content_type == ContentType.PHOTO or (message.content_type == ContentType.DOCUMENT and message.document.mime_type.startswith('image')):
            user_photos = photos_id.get(user_id, [])
            file_id = message.photo[-1].file_id if message.content_type == ContentType.PHOTO else message.document.file_id
            user_photos.append(file_id)
            photos_id[user_id] = user_photos
            kb = types.InlineKeyboardMarkup()
            kb.add(types.InlineKeyboardButton('Конвертировать в PDF', callback_data='convert_pdf'))
            kb.add(types.InlineKeyboardButton('Конвертировать в DOCX', callback_data='convert_docx'))
            await message.answer(f'Изображение добавлено! Всего изображений: {len(user_photos)}\n\nНажмите кнопку для конвертации.', reply_markup=kb)
        else:
            await message.answer('В режиме конвертации изображений принимаются только изображения. Пожалуйста, отправьте изображение.')
    elif mode == 'mode_2':
        if message.content_type == ContentType.DOCUMENT and message.document.mime_type in ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            user_docs = docs_id.get(user_id, [])
            user_docs.append(message.document.file_id)
            docs_id[user_id] = user_docs
            kb = types.InlineKeyboardMarkup()
            kb.add(types.InlineKeyboardButton('Объединить PDF', callback_data='merge_pdf'))
            kb.add(types.InlineKeyboardButton('Объединить DOCX', callback_data='merge_docx'))
            await message.answer(f'Документ добавлен! Всего документов: {len(user_docs)}\n\nНажмите кнопку для объединения.', reply_markup=kb)
        else:
            await message.answer('В режиме объединения принимаются только PDF или DOCX файлы. Пожалуйста, отправьте файл.')
    else:
        await message.answer('Пожалуйста, сначала выберите режим работы бота командой /start.')

# Конвертация изображений
@dp.callback_query_handler(lambda c: c.data in ['convert_pdf', 'convert_docx'])
async def convert_images_handler(query: types.CallbackQuery):
    user_id = str(query.message.chat.id)
    await query.answer()
    file_names[user_id] = None
    bot_data_key = 'convert_format_' + user_id
    bot[bot_data_key] = query.data
    await query.message.answer('Введите желаемое имя файла (без расширения):')
    await ConvertState.waiting_for_filename.set()

@dp.message_handler(state=ConvertState.waiting_for_filename, content_types=types.ContentTypes.TEXT)
async def process_convert_filename(message: types.Message, state: FSMContext):
    user_id = str(message.chat.id)
    filename = message.text.strip()
    file_names[user_id] = filename

    convert_format = bot.get('convert_format_' + user_id, 'convert_pdf')
    try:
        await convert_and_send_file(user_id, filename, convert_format)
    except Exception as e:
        await message.answer(f'Ошибка при конвертации: {e}')
        await state.finish()
        return

    await message.answer(f"Файл '{filename}' успешно создан и отправлен.", reply_markup=back_keyboard())
    photos_id.pop(user_id, None)
    docs_id.pop(user_id, None)
    file_names.pop(user_id, None)
    user_modes[user_id] = None
    await state.finish()

async def convert_and_send_file(user_id, filename, file_format):
    user_dir = os.path.join(dir_path, 'UserData', user_id)
    os.makedirs(user_dir, exist_ok=True)

    temp_photos = photos_id.get(user_id, [])
    if not temp_photos:
        raise IndexError("Нет изображений для конвертации")

    for idx, file_id in enumerate(temp_photos, start=1):
        path = os.path.join(user_dir, f'{idx}.jpg')
        await bot.download_file_by_id(file_id, path)

    if file_format == 'convert_pdf':
        images = []
        image_files = sorted([f for f in os.listdir(user_dir) if f.endswith('.jpg')],
                             key=lambda x: int(os.path.splitext(x)[0]))
        for im_name in image_files:
            im = Image.open(os.path.join(user_dir, im_name))
            images.append(im.convert('RGB'))
        pdf_path = os.path.join(user_dir, f'{filename}.pdf')
        if images:
            images[0].save(pdf_path, save_all=True, append_images=images[1:])
        upload_path = pdf_path

    elif file_format == 'convert_docx':
        doc = Document()
        image_files = sorted([f for f in os.listdir(user_dir) if f.endswith('.jpg')],
                             key=lambda x: int(os.path.splitext(x)[0]))
        for im_name in image_files:
            img_path = os.path.join(user_dir, im_name)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_page_break()
        docx_path = os.path.join(user_dir, f'{filename}.docx')
        doc.save(docx_path)
        upload_path = docx_path

    else:
        raise ValueError("Неизвестный формат конвертации")

    await bot.send_document(user_id, types.InputFile(upload_path))

# Обработка нажатий на кнопки объединения
@dp.callback_query_handler(lambda c: c.data in ['merge_pdf', 'merge_docx'])
async def merge_files_handler(query: types.CallbackQuery):
    user_id = str(query.message.chat.id)
    await query.answer()
    file_names[user_id] = None
    bot_data_key = 'merge_format_' + user_id
    bot[bot_data_key] = query.data
    await query.message.answer('Введите желаемое имя итогового файла (без расширения):')
    await MergeState.waiting_for_filename.set()

@dp.message_handler(state=MergeState.waiting_for_filename, content_types=types.ContentTypes.TEXT)
async def process_merge_filename(message: types.Message, state: FSMContext):
    user_id = str(message.chat.id)
    filename = message.text.strip()

    merge_format = bot.get('merge_format_' + user_id)
    user_docs = docs_id.get(user_id, [])
    if not user_docs:
        await message.answer('Нет файлов для объединения. Пожалуйста, отправьте PDF или DOCX файлы.')
        await state.finish()
        return

    user_dir = os.path.join(dir_path, 'UserData', user_id)
    os.makedirs(user_dir, exist_ok=True)

    # Скачиваем все файлы
    local_files = []
    for idx, file_id in enumerate(user_docs, start=1):
        file_ext = '.pdf' if merge_format == 'merge_pdf' else '.docx'
        path = os.path.join(user_dir, f'{idx}{file_ext}')
        await bot.download_file_by_id(file_id, path)
        local_files.append(path)

    try:
        if merge_format == 'merge_pdf':
            merger = PdfMerger()
            for pdf_file in local_files:
                merger.append(pdf_file)
            output_path = os.path.join(user_dir, f'{filename}.pdf')
            with open(output_path, 'wb') as f_out:
                merger.write(f_out)
            merger.close()
        elif merge_format == 'merge_docx':
            merged_document = Document(local_files[0])
            for docx_file in local_files[1:]:
                sub_doc = Document(docx_file)
                for element in sub_doc.element.body:
                    merged_document.element.body.append(element)
            output_path = os.path.join(user_dir, f'{filename}.docx')
            merged_document.save(output_path)
        else:
            await message.answer('Неизвестный формат для объединения.')
            await state.finish()
            return

        await message.answer(f"Файл '{filename}' успешно создан и отправлен.", reply_markup=back_keyboard())
        await bot.send_document(user_id, types.InputFile(output_path))

    except Exception as e:
        await message.answer(f'Ошибка при объединении файлов: {e}')

    docs_id.pop(user_id, None)
    file_names.pop(user_id, None)
    user_modes[user_id] = None
    await state.finish()

if __name__ == '__main__':
    executor.start_polling(dp, skip_updates=True)
